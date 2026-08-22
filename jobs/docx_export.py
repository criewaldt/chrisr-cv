"""Render a tailored resume as .docx.

Worth having alongside the PDF: several ATS parse Word documents more reliably
than PDF, and a few reject PDF outright. Section order mirrors resume/pdf.py so the
two formats stay in step.
"""
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x43, 0x4E, 0x5E)
MUTED = RGBColor(0x6C, 0x75, 0x7D)

SKILL_CATEGORIES = [('web', 'Web Development'), ('language', 'Languages'),
                    ('cloud', 'Cloud & DevOps'), ('ai', 'AI & ML'),
                    ('methodology', 'Methodologies')]


def _blocks(html):
    """Flatten stored HTML into (text, is_bullet) pairs."""
    if not html:
        return []
    soup = BeautifulSoup(html, 'html.parser')
    out = []
    for node in soup.children:
        name = getattr(node, 'name', None)
        if name in ('ul', 'ol'):
            for item in node.find_all('li', recursive=False):
                text = ' '.join(item.get_text(' ', strip=True).split())
                if text:
                    out.append((text, True))
        else:
            text = ' '.join(node.get_text(' ', strip=True).split()) if name else str(node).strip()
            if text:
                out.append((text, False))
    return out


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    return p


def _date_range(job):
    start = job.start_date.strftime('%b %Y') if job.start_date else ''
    end = 'Present' if (job.is_current or not job.end_date) else job.end_date.strftime('%b %Y')
    span = ' - '.join(p for p in (start, end) if p)
    return ' | '.join(p for p in (span, job.location) if p)


def render_resume_docx(view, profile_links=()):
    """``view`` is a TailoredResumeView -- the same object the PDF renderer takes."""
    doc = Document()
    section = doc.sections[0]
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(section, attr, Inches(0.6))
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(view.name)
    run.bold = True
    run.font.size = Pt(20)

    title = view.desired_title or view.current_title
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.size = Pt(11.5)
        r.font.color.rgb = ACCENT

    contact = ' | '.join(x for x in (view.email, view.phone,
                                     *(label for label, _ in profile_links)) if x)
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(contact)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    summary = _blocks(view.professional_summary.summary_html)
    if summary:
        _heading(doc, 'Career Summary')
        for text, bullet in summary:
            doc.add_paragraph(text, style='List Bullet' if bullet else None)

    jobs = sorted(view.employment_history.all(), key=lambda j: j.sort_order)
    if jobs:
        _heading(doc, 'Work Experience')
        for job in jobs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(job.job_title or '')
            r.bold = True
            r.font.size = Pt(10.5)
            if job.company_name:
                c = p.add_run(f'  —  {job.company_name}')
                c.bold = True
                c.font.color.rgb = ACCENT
            meta = doc.add_paragraph()
            meta.paragraph_format.space_after = Pt(2)
            m = meta.add_run(_date_range(job))
            m.italic = True
            m.font.size = Pt(8.5)
            m.font.color.rgb = MUTED
            for text, bullet in _blocks(job.description_html):
                doc.add_paragraph(text, style='List Bullet' if bullet else None)

    keywords = list(view.keywords.all())
    if keywords:
        _heading(doc, 'Skills & Tools')
        grouped = {}
        for k in keywords:
            grouped.setdefault((k.category or '').strip().lower(), []).append(k.name)
        ordered = [(key, label) for key, label in SKILL_CATEGORIES if key in grouped]
        known = {key for key, _ in SKILL_CATEGORIES}
        ordered += [(k, (k or 'Other').replace('_', ' ').title())
                    for k in sorted(grouped) if k not in known]
        for key, label in ordered:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(f'{label}: ')
            r.bold = True
            p.add_run(', '.join(grouped[key]))

    education = list(view.education.all())
    if education:
        _heading(doc, 'Education')
        for edu in education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.add_run(edu.degree or '').bold = True
            meta = ' | '.join(x for x in (edu.school, edu.time) if x)
            if meta:
                r = p.add_run(f'  {meta}')
                r.font.size = Pt(9)
                r.font.color.rgb = MUTED

    awards = list(view.awards.all())
    if awards:
        _heading(doc, 'Licenses & Awards')
        for award in awards:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.add_run(award.name or '').bold = True
            if award.description:
                r = p.add_run(f'  {award.description}')
                r.font.size = Pt(9)
                r.font.color.rgb = MUTED

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
